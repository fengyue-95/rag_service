"""文件内容提取工具模块

支持提取以下格式的文本内容：
- txt, md (纯文本)
- pdf (PDF文档)
- doc, docx (Word文档)
- xls, xlsx (Excel文档)
"""

import re
from pathlib import Path
from typing import Optional


def extract_text_from_file(file_path: Path) -> Optional[str]:
    """
    从文件中提取文本内容

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容，提取失败返回 None
    """
    file_ext = file_path.suffix.lower()
    
    if file_ext in {'.txt', '.md'}:
        return _extract_text_from_txt(file_path)
    elif file_ext == '.pdf':
        return _extract_text_from_pdf(file_path)
    elif file_ext in {'.doc', '.docx'}:
        return _extract_text_from_doc(file_path)
    elif file_ext in {'.xls', '.xlsx'}:
        return _extract_text_from_excel(file_path)
    else:
        print(f"不支持的文件类型: {file_ext}")
        return None


def clean_text(content: str) -> str:
    """
    清洗文本中的乱码和无效字符

    Args:
        content: 原始文本

    Returns:
        清洗后的文本
    """
    if not content:
        return content
    
    # 替换常见的乱码字符（ Replacement Character U+FFFD）
    content = content.replace('\ufffd', '')
    
    # 移除控制字符（保留换行符\t）
    content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', content)
    
    # 替换多个连续空白字符为单个空格（保留换行）
    content = re.sub(r'[ \t]+', ' ', content)
    
    # 移除每行开头和结尾的空白
    lines = [line.strip() for line in content.split('\n')]
    # 过滤空行
    lines = [line for line in lines if line]
    content = '\n'.join(lines)
    
    return content


def _extract_text_from_txt(file_path: Path) -> Optional[str]:
    """从 txt/md 文件提取文本"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'iso-8859-1', 'latin-1']
    
    for encoding in encodings:
        try:
            with file_path.open('r', encoding=encoding) as f:
                content = f.read()
            print(f"使用 {encoding} 编码成功读取文件: {file_path.name}")
            content = clean_text(content)
            print(f"内容预览:\n{content[:1000]}...")
            return content
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"使用 {encoding} 编码读取失败: {str(e)}")
            continue
    
    print(f"无法读取文件内容: {file_path.name}")
    return None


def _extract_text_from_pdf(file_path: Path) -> Optional[str]:
    """从 PDF 文件提取文本（含表格识别）"""
    text_content = []
    table_content = []
    
    try:
        import PyPDF2
        with file_path.open('rb') as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                # 提取普通文本
                text = page.extract_text()
                if text:
                    text_content.append(text)
                
                # 尝试使用 pdfplumber 提取表格（发票等表格文档）
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        if i < len(pdf.pages):
                            page_obj = pdf.pages[i]
                            tables = page_obj.extract_tables()
                            if tables:
                                for table in tables:
                                    table_str = _format_table(table)
                                    if table_str:
                                        table_content.append(f"[表格区域]\n{table_str}")
                                        print(f"第 {i+1} 页发现表格")
                    pdfplumber_available = True
                except ImportError:
                    pdfplumber_available = False
                except Exception as e:
                    pdfplumber_available = False
                    print(f"pdfplumber 提取表格失败（{file_path.name}）: {str(e)}")
            
            # 合并文本和表格内容
            all_content = []
            if text_content:
                all_content.append('\n'.join(text_content))
            if table_content:
                all_content.append('\n\n'.join(table_content))
            
            content = '\n\n'.join(all_content)
            content = clean_text(content)
            
            if pdfplumber_available:
                print(f"成功提取 PDF 文本内容（含表格）: {file_path.name}, 共 {len(reader.pages)} 页")
            else:
                print(f"成功提取 PDF 文本内容: {file_path.name}, 共 {len(reader.pages)} 页")
                print("提示：安装 pdfplumber 可支持表格识别（pip install pdfplumber）")
            
            print(f"内容预览:\n{content[:1000]}...")
            return content
            
    except ImportError:
        print("PyPDF2 未安装，无法提取 PDF 内容")
        return None
    except Exception as e:
        print(f"提取 PDF 内容失败: {str(e)}")
        return None


def _format_table(table: list) -> str:
    """格式化表格内容为易读文本"""
    if not table or len(table) < 1:
        return ""
    
    lines = []
    for row in table:
        if row:
            # 清理每行中的单元格
            cleaned_row = []
            for cell in row:
                if cell is None:
                    cleaned_row.append("")
                else:
                    cell_str = str(cell).strip()
                    cell_str = re.sub(r'\s+', ' ', cell_str)
                    cleaned_row.append(cell_str)
            lines.append(' | '.join(cleaned_row))
    
    return '\n'.join(lines)


def _extract_text_from_doc(file_path: Path) -> Optional[str]:
    """从 Word 文件提取文本（含表格）"""
    text_content = []
    table_content = []
    
    try:
        from docx import Document
        doc = Document(file_path)
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # 提取表格内容（发票、报表等）
        for i, table in enumerate(doc.tables):
            table_str = _format_docx_table(table)
            if table_str:
                table_content.append(f"[表格区域 {i+1}]\n{table_str}")
                print(f"发现 Word 表格 {i+1}")
        
        # 合并文本和表格内容
        all_content = []
        if text_content:
            all_content.append('\n'.join(text_content))
        if table_content:
            all_content.append('\n\n'.join(table_content))
        
        content = '\n\n'.join(all_content)
        content = clean_text(content)
        
        print(f"成功提取 Word 文本内容: {file_path.name}")
        if table_content:
            print(f"包含 {len(table_content)} 个表格")
        print(f"内容预览:\n{content[:1000]}...")
        return content
        
    except ImportError:
        print("python-docx 未安装，无法提取 Word 内容")
        return None
    except Exception as e:
        print(f"提取 Word 内容失败: {str(e)}")
        return None


def _format_docx_table(table) -> str:
    """格式化 Word 表格内容"""
    lines = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            cell_text = re.sub(r'\s+', ' ', cell_text)
            row_data.append(cell_text)
        if any(row_data):  # 只添加非空行
            lines.append(' | '.join(row_data))
    return '\n'.join(lines)


def _extract_text_from_excel(file_path: Path) -> Optional[str]:
    """从 Excel 文件提取文本（含表格）"""
    try:
        import pandas as pd
        
        # 读取所有工作表
        excel_file = pd.ExcelFile(file_path)
        all_content = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            
            # 添加工作表名称
            all_content.append(f"[工作表: {sheet_name}]")
            
            # 转换为文本格式
            # 处理表格头
            if not df.empty:
                headers = df.columns.tolist()
                header_line = ' | '.join(str(h) for h in headers)
                all_content.append(header_line)
                
                # 处理数据行
                for _, row in df.iterrows():
                    row_data = []
                    for val in row:
                        if pd.isna(val):
                            row_data.append("")
                        elif isinstance(val, float):
                            row_data.append(f"{val:.2f}")
                        else:
                            row_data.append(str(val))
                    all_content.append(' | '.join(row_data))
            
            all_content.append("")  # 空行分隔
        
        content = '\n'.join(all_content)
        content = clean_text(content)
        
        print(f"成功提取 Excel 内容: {file_path.name}, 共 {len(excel_file.sheet_names)} 个工作表")
        print(f"内容预览:\n{content[:1000]}...")
        return content
        
    except ImportError:
        print("pandas 未安装，无法提取 Excel 内容")
        return None
    except Exception as e:
        print(f"提取 Excel 内容失败: {str(e)}")
        return None


def get_supported_extensions() -> set:
    """获取支持的文件扩展名"""
    return {'.txt', '.md', '.pdf', '.doc', '.docx', '.xls', '.xlsx'}
